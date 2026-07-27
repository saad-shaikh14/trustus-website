# -*- coding: utf-8 -*-
"""Generates the TRUSTUS Portal HR Guide docx for Sajid — one-time deliverable,
not a repeated-generation tool, but kept as a script (not hand-edited XML) so
it can be regenerated if the workflow changes. Run: python build_hr_guide.py
"""
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

NAVY = RGBColor(0x1C, 0x2B, 0x4A)
PINK = RGBColor(0xC8, 0x52, 0x6A)
GREEN = RGBColor(0x3D, 0x6A, 0x5A)
GRAY = RGBColor(0x44, 0x44, 0x44)
LIGHTGRAY = RGBColor(0x88, 0x88, 0x88)

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LOGO = os.path.join(ROOT, "images", "logo.png")
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "TRUSTUS_Portal_HR_Guide.docx")

doc = docx.Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = GRAY

def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.size = Pt(17)
    r.font.bold = True
    r.font.color.rgb = NAVY
    # underline rule
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "C8526A")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = PINK
    return p

def body(text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.font.bold = True
        r.font.color.rgb = NAVY
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.font.bold = True
        r.font.color.rgb = NAVY
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.font.bold = True
        r.font.color.rgb = NAVY
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def note(text, label="NOTE"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run(f"{label}: ")
    r.font.bold = True
    r.font.color.rgb = PINK
    r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.italic = True
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GRAY
    return p

def link_line(label, url):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(f"{label}: ")
    r.font.bold = True
    r.font.color.rgb = NAVY
    r2 = p.add_run(url)
    r2.font.color.rgb = PINK
    r2.font.underline = True
    return p

# ---------- Cover ----------
if os.path.exists(LOGO):
    doc.add_picture(LOGO, width=Inches(1.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(20)
r = title.add_run("TRUSTUS Care — Job Portal Guide")
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(24)
r = sub.add_run("A complete, start-to-finish guide to reviewing candidates, sending invites,\nrunning the competency test, and everything else in the HR portal.")
r.font.size = Pt(12)
r.font.italic = True
r.font.color.rgb = LIGHTGRAY

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = date_p.add_run("Prepared for Sajid — 27 July 2026")
r.font.size = Pt(10)
r.font.color.rgb = LIGHTGRAY

doc.add_page_break()

# ---------- 1. Getting started ----------
h1("1. Getting Started")

h2("1.1 The Portal Address")
link_line("Portal URL", "https://trustuscare.com/portal")
body("This is where you log in to review candidates, send invites, run the competency test, and record interview outcomes. Bookmark this page.")

h2("1.2 Logging In")
numbered("Go to the portal address above.")
numbered("Enter your login email: info@trustuscare.com")
numbered("Enter your password (the one already set up for you).")
numbered("Click Sign In.")
note("If you've forgotten the password or it's not working, contact Saad — there is no self-service \"forgot password\" link on this portal, so it needs to be reset from the admin side.")

h2("1.3 The Careers Page (where candidates apply)")
link_line("Careers page URL", "https://trustuscare.com/careers.html")
body("This is the public page candidates use to apply. They pick a role (Care Coordinator, Field Care Worker, or Learning Disability Support Worker), fill in their name, phone, email, and upload their CV. You don't need to do anything on this page yourself — it's just useful to know what candidates see before they reach you.")

# ---------- 2. Dashboard ----------
h1("2. The Dashboard")
body("Once logged in, you'll see the main dashboard:")
bullet("A row of numbers at the top showing Total applications, Form Complete, and Interview Invited counts.")
bullet("A search box — search by candidate name or email.")
bullet("A filter dropdown — narrow the list down to a specific stage.")
bullet("A list of every candidate on the left; click any candidate to open their full record on the right.")

note("The filter dropdown also lists \"Test Complete\" and \"Shortlisted\" as options. These aren't used in the current workflow — clicking them will just show an empty list. The stages that are actually used are covered in the next section.")

h2("2.1 Candidate Stages (Statuses)")
bullet("A candidate has just submitted their CV and hasn't completed the full application form yet.", "CV Received — ")
bullet("The candidate has completed the full application form (personal details, address history, employment history, education, referees). Ready for your review.", "Form Complete — ")
bullet("You've sent them an interview invite with a date, time, and location.", "Interview Invited — ")
bullet("You've marked this candidate as not progressing. This cannot be undone from the portal — see Section 8.", "Not Progressing (Rejected) — ")

# ---------- 3. New application ----------
h1("3. Step by Step: A New Application Comes In")
numbered("A candidate applies via the careers page and uploads their CV.")
numbered("You automatically receive an email titled \"New CV Received\" at info@trustuscare.com, with the candidate's name, role, email, and phone.")
numbered("The candidate automatically receives an email asking them to complete the full application form. This link is valid for 7 days. You don't need to do anything at this stage.")
note("Everything up to this point happens automatically — no action needed from you.")

# ---------- 4. Form complete ----------
h1("4. Step by Step: Candidate Completes the Application Form")
numbered("You receive an email titled \"Form 1 Complete\".")
numbered("The candidate's status in the portal changes to Form Complete.")
numbered("You can now open their record and review everything.")

# ---------- 5. Reviewing a candidate ----------
h1("5. Reviewing a Candidate")
body("Click on a candidate in the list to open their full record. You'll see:")
bullet("Email and phone number (click either to email or call directly).")
bullet("A Download CV button.")
bullet("Their full application form: address history, employment history, education, referees, and declaration.")
bullet("A timeline showing exactly what's happened so far and when (see Section 9).")
bullet("An HR Notes box for your own private notes on this candidate — type your notes and click Save Notes.")

note("Downloading the CV automatically opens the \"Send Interview Invite\" box a moment later, as a handy reminder in case you're ready to move them forward.")

# ---------- 6. Interview invite ----------
h1("6. Sending an Interview Invite")
numbered("Open the candidate's record and click Send Interview Invite.")
numbered("Fill in the Interview Date, Interview Time, Location/Address, and any Additional Notes (e.g. \"please bring proof of ID and right to work documents\").")
numbered("Click Send Invite.")
numbered("The candidate receives an email with all these details.")
numbered("Their status changes to Interview Invited.")

note("This invite email does NOT include a link to the competency test — that's a separate step, sent separately. See Section 7 below.", label="IMPORTANT")

# ---------- 7. Competency test ----------
h1("7. The Competency Test")

h2("7.1 What It Is")
body("A 20-question multiple-choice test covering five areas: Core Values & Independence, Safeguarding & Rules, Medication & Transfers, Emergency Situations, and Communication. A candidate needs 12 out of 20 to pass. The candidate never sees their own score — only you do.")

h2("7.2 Sending the Test to a Candidate")
numbered("Open the candidate's record and click Generate Competency Test Link.")
numbered("A short message explains a link will be copied to your clipboard.")
numbered("Click Generate & Copy Link.")
numbered("The link is now on your clipboard — paste it into an email, text message, or WhatsApp to the candidate yourself.")
note("You have to send this link to the candidate yourself — the system does not email it automatically. This is different from the interview invite, which does send automatically.", label="IMPORTANT")

h2("7.3 What Happens When They Take It")
bullet("The candidate answers all 20 questions and submits.")
bullet("You automatically receive an email showing PASS or FAIL, their score out of 20, and a breakdown by section.")
bullet("If they accidentally close the browser tab before submitting, you can just generate the same link again — nothing is lost, since no score exists yet.")

h2("7.4 Viewing Their Answers")
body("Once a candidate has completed the test, a View Test Results button appears on their record. Click it to see their score, pass/fail, and every question with their chosen answer marked correct or incorrect.")

# ---------- 8. Second chance ----------
h1("8. Giving a Candidate a Second Chance")
body("If you think a candidate deserves another attempt at the test — maybe they had a bad day, or you believe they're a strong fit despite a weak first score — you can let them retake it, without losing the original result.")
numbered("Open the candidate's record. Since they already have a result, the button now reads Give Second Chance (Regenerate Test).")
numbered("Click it. The message explains: their current score will be kept on file as a previous attempt, and generating a new link lets them retake the test — the new result becomes the official one.")
numbered("Click Generate & Copy Link, exactly as before, and send it to the candidate yourself.")
numbered("Once they submit again, their new score becomes the one shown on their record — but their original attempt is not deleted.")
body("Under the score card, you'll see a \"Previous Attempts\" list showing every earlier attempt with its date and score, so nothing is ever lost.")
note("You can repeat this as many times as you genuinely need to — there's no limit on the number of attempts.")

# ---------- 9. Scorecard ----------
h1("9. The Interviewer Scorecard")
body("Separate from the multiple-choice test, this is where you (or whoever conducts the interview) record your own assessment of the candidate. It appears as a Scorecard tab once a candidate has reached Form Complete or later.")
bullet("Core Values & Independence")
bullet("Safeguarding & Rules")
bullet("Medication & Transfers")
bullet("Specialized Communication")
bullet("Professional Boundaries")
body("Rate each of the five domains above out of 5 stars, then fill in:")
bullet("Strengths (free text)")
bullet("Development Areas (free text)")
bullet("Outcome — PASS, HOLD, or FAIL")
bullet("Interviewer Name")
body("Click Save when done.")
note("The scorecard is about the candidate overall, not tied to one specific test attempt — if you give a candidate a second chance on the test (Section 8), their scorecard stays exactly as you left it.")

# ---------- 10. Rejecting ----------
h1("10. Marking a Candidate as Not Progressing")
numbered("Open the candidate's record and click the red ✖ Not Progressing button.")
numbered("Confirm when asked.")
numbered("Their status changes to Rejected.")
note("This cannot be undone from the portal — there is no \"undo\" button by design, to avoid accidental status changes. If you click this by mistake, contact Saad to have it corrected directly.", label="IMPORTANT")

# ---------- 11. Timeline ----------
h1("11. The Candidate Timeline")
body("Every candidate's record shows a timeline of exactly what's happened and when:")
bullet("CV Submitted")
bullet("Application Form Completed")
bullet("CV Downloaded by HR")
bullet("Interview Invite Sent")
bullet("Competency Test Submitted")
bullet("Test Results Viewed by HR")
body("Each shows a green tick and a timestamp once it's happened, or \"Not yet\" if it hasn't.")

# ---------- 12. Emails ----------
h1("12. Understanding Your Emails")
body("Everything below arrives at info@trustuscare.com automatically. Here's what each one means:")

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Email Subject"
hdr[1].text = "What It Means"
for c in hdr:
    set_cell_shading(c, "1C2B4A")
    for p in c.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.bold = True

rows_data = [
    ("New CV Received — [Name] ([Role])", "A new candidate has applied. No action needed yet — they'll get the full application form automatically."),
    ("Form 1 Complete — [Name] ([Role])", "The candidate has finished the full application form. Time to review them in the portal."),
    ("Competency Test PASS/FAIL — [Name] ([Score]/20)", "The candidate has completed the multiple-choice test. If it says \"Attempt 2\" or higher, this is a retake you set up (Section 8)."),
    ("TRUSTUS Portal Health Check Failed", "An automated daily check found the portal's database unreachable. This is NOT a normal notification — forward it to Saad immediately, as candidates may not be able to apply until it's resolved."),
]
for subj, meaning in rows_data:
    row = table.add_row().cells
    row[0].text = subj
    row[1].text = meaning
    for c in row:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ---------- 13. Help ----------
h1("13. If Something Goes Wrong")
bullet("Contact Saad — there's no self-service reset.", "Forgotten password: ")
bullet("Wait a minute and refresh; if it persists, contact Saad. There is now an automated check that emails info@trustuscare.com if the system goes down (see Section 12) — if you get that email, forward it to Saad straight away.", "Portal won't load: ")
bullet("If you clicked ✖ Not Progressing by mistake, contact Saad to have the status corrected directly — this cannot be undone from the portal itself.", "Rejected the wrong candidate: ")
bullet("If a candidate emails saying they applied but never heard back, check the portal search first — if they genuinely don't appear anywhere in the list, contact Saad, since this could indicate a technical issue worth investigating.", "Candidate says they applied but you can't find them: ")

# ---------- 14. Data & privacy ----------
h1("14. A Note on Data & Privacy")
body("Candidate personal details (National Insurance number, date of birth, address history) are shown as read-only in the portal — they can't be edited from here. There is also no delete button anywhere in the portal; this is intentional, to protect candidate data. If a candidate ever asks for their data to be deleted, contact Saad rather than trying to remove anything yourself.")

doc.add_paragraph().paragraph_format.space_after = Pt(20)
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("TRUSTUS Care · Office G24, 47 Clarendon Road, Watford, WD17 1HP · 020 3411 1218 · info@trustuscare.com")
r.font.size = Pt(9)
r.font.color.rgb = LIGHTGRAY

doc.save(OUT)
print(f"Saved: {OUT}")
